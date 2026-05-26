"""文档结构解析器 — 三段式检测流水线"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import utils
from config import (
    FONT_RULES, PAGE_MARGINS, LINE_SPACING, PARA_SPACING_BEFORE, PARA_SPACING_AFTER,
)

# 对齐方式值 → 可读名称
_ALIGN_NAMES = {
    None: None,
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE",
}


def parse_document(filepath: str) -> dict:
    """
    解析 .docx 文件，返回结构化分析结果。

    返回格式:
    {
        "paragraphs": [
            {
                "index": 0,
                "text": "...",
                "element_type": "title",
                "issues": [{"desc": "...", "severity": "error/warning"}],
                "current_format": {...},
                "expected_format": {...}
            },
            ...
        ],
        "issue_count": 3
    }
    """
    doc = Document(filepath)
    paragraphs = []
    total = len(doc.paragraphs)

    # 找到真正的公文题目起始位置（跳过开头的附件行）
    title_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if utils.ATTACHMENT.match(text):
            continue  # 跳过开头附件行，找后面的真正标题
        title_idx = i
        break

    # ── 第一段：逐段初步分类 ──
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        is_first = (i == title_idx)
        etype = utils.classify_paragraph(text, i, total, is_first)

        info = {
            "index": i,
            "text": text.strip(),
            "element_type": etype,
            "issues": [],
            "current_format": _capture_format(para),
            "expected_format": None,
        }
        paragraphs.append(info)

    # ── 第二段：上下文精修 ──
    _refine_context(paragraphs)

    # ── 第三段：逐段格式合规检查 ──
    subtitle_bold_set = _detect_subtitle_bold(paragraphs)
    issue_count = 0
    for i, info in enumerate(paragraphs):
        _check_format(info, i in subtitle_bold_set)
        issue_count += len(info["issues"])

    # ── 第四段：文档级格式检查（页边距等）──
    page_issues = _check_page_setup(doc)

    return {
        "paragraphs": paragraphs,
        "issue_count": issue_count + len(page_issues),
        "page_issues": page_issues,
    }


def _refine_context(paragraphs: list):
    """从文档末尾向前扫描，识别附件/附注/落款/抄送"""
    n = len(paragraphs)

    for i in range(n - 1, -1, -1):
        text = paragraphs[i]["text"]
        etype = paragraphs[i]["element_type"]

        if etype == "empty":
            continue

        # 日期
        if utils.DATE.match(text):
            paragraphs[i]["element_type"] = "signature_date"
            continue

        # 抄送
        if utils.CC.match(text):
            paragraphs[i]["element_type"] = "cc"
            continue

        # 附注
        if utils.NOTE.match(text):
            paragraphs[i]["element_type"] = "note"
            continue

        # 附件标题（以"附件"开头且较短）
        if utils.ATTACHMENT.match(text) and len(text) < 50:
            paragraphs[i]["element_type"] = "attachment_title"
            continue

    # 正向传播：附件标题后的 body 段落标记为附件内容
    for i in range(1, n):
        if paragraphs[i]["element_type"] == "body":
            prev_type = paragraphs[i - 1]["element_type"]
            if prev_type in ("attachment_title", "attachment"):
                paragraphs[i]["element_type"] = "attachment"

    # 第三遍：落款检测 — 文档末尾附近的短文本
    signature_start = None
    for i in range(n - 1, -1, -1):
        if paragraphs[i]["element_type"] == "signature_date":
            signature_start = i
            break

    if signature_start:
        for i in range(signature_start - 1, max(signature_start - 2, -1), -1):
            etype = paragraphs[i]["element_type"]
            text = paragraphs[i]["text"]
            if etype == "body" and len(text) < 20:
                paragraphs[i]["element_type"] = "signature"
            elif etype not in ("empty", "body", "signature", "speaker", "figure_caption"):
                break


def _capture_format(para) -> dict:
    """捕获段落当前格式（含行距）"""
    fmt = {}
    pf = para.paragraph_format
    fmt["alignment"] = str(pf.alignment) if pf.alignment is not None else None
    indent = pf.first_line_indent
    fmt["first_line_indent"] = f"{indent / 360000:.2f}cm" if indent else "无"

    # 行距（从 XML 读取）
    fmt["line_spacing"] = "?"
    fmt["line_rule"] = "?"
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            lv = sp.get(qn('w:line'))
            lr = sp.get(qn('w:lineRule'))
            if lv and lr:
                if lr == 'exact':
                    fmt["line_spacing"] = f"{int(lv)/20:.0f}磅(固定)"
                elif lr == 'auto':
                    fmt["line_spacing"] = f"{int(lv)/240:.1f}倍"
                else:
                    fmt["line_spacing"] = f"{int(lv)/20:.0f}磅"
                fmt["line_rule"] = lr

    # 段前段后间距 — 复用 pPr，读 w:before/w:after（twips），兜底 python-docx
    sb_val = 0
    sa_val = 0
    if pPr is not None:
        sp_ps = pPr.find(qn('w:spacing'))
        if sp_ps is not None:
            sb_raw = sp_ps.get(qn('w:before'))
            sa_raw = sp_ps.get(qn('w:after'))
            if sb_raw: sb_val = int(sb_raw)
            if sa_raw: sa_val = int(sa_raw)
    if sb_val == 0:
        sb_raw = pf.space_before
        sb_val = int(sb_raw) if sb_raw else 0
    if sa_val == 0:
        sa_raw = pf.space_after
        sa_val = int(sa_raw) if sa_raw else 0
    # 样式继承兜底
    if sb_val == 0 or sa_val == 0:
        try:
            style = para.style
            if style and style.paragraph_format:
                spf = style.paragraph_format
                if sb_val == 0 and spf.space_before:
                    sb_val = int(spf.space_before)
                if sa_val == 0 and spf.space_after:
                    sa_val = int(spf.space_after)
        except Exception:
            pass
    if abs(sb_val) > 1000:
        fmt["space_before"] = f"{sb_val / 12700:.0f}pt"
    else:
        fmt["space_before"] = f"{sb_val / 20:.0f}pt" if sb_val else "0pt"
    if abs(sa_val) > 1000:
        fmt["space_after"] = f"{sa_val / 12700:.0f}pt"
    else:
        fmt["space_after"] = f"{sa_val / 20:.0f}pt" if sa_val else "0pt"

    # 首句加粗检测：第一个句号前的 run 是否有加粗
    fmt["first_sentence_bold"] = False
    if para.runs:
        first_dot = para.text.find('。') if para.text else -1
        pos = 0
        for run in para.runs:
            t = run.text if run.text else ''
            if first_dot >= 0 and pos < first_dot + 1:
                if run.font.bold:
                    fmt["first_sentence_bold"] = True
            pos += len(t)

    if para.runs:
        run = para.runs[0]
        fmt["font_name"] = run.font.name
        fmt["font_size"] = str(run.font.size) if run.font.size else None
        fmt["bold"] = run.font.bold
    else:
        fmt["font_name"] = None
        fmt["font_size"] = None
        fmt["bold"] = None

    return fmt


def _detect_subtitle_bold(paragraphs: list) -> set:
    """检测哪些段落应当有小标题加粗，返回段落 index 集合"""
    import re
    should_bold = set()

    # ── 模式1：连续3段以上首句等长 ──
    first_sentences = []
    for info in paragraphs:
        text = info["text"]
        dot = text.find('。')
        if dot == -1:
            first_sentences.append(None)
        else:
            first_sentences.append((info["index"], len(text[:dot])))

    i = 0
    while i < len(first_sentences):
        if first_sentences[i] is None:
            i += 1
            continue
        _, length = first_sentences[i]
        j = i + 1
        while j < len(first_sentences):
            if first_sentences[j] is None or first_sentences[j][1] != length:
                break
            j += 1
        if j - i >= 3:
            for k in range(i, j):
                should_bold.add(first_sentences[k][0])
            i = j
        else:
            i += 1

    # ── 模式2：段内"一是/二是/三是"等分句 ≥3 且等长 ──
    branch_re = re.compile(r'[一二三四五六七八九十]+[、，,是]|\d+[、，,\.是]')
    for info in paragraphs:
        text = info["text"]
        markers = []
        for m in branch_re.finditer(text):
            pos = m.start()
            if pos == 0 or (pos > 0 and text[pos-1] in '。；;，,、\n'):
                markers.append(m.start())
        if len(markers) < 3:
            continue
        seg_lens = []
        for idx, seg_start in enumerate(markers):
            dot = text.find('。', seg_start)
            seg_end = dot + 1 if dot != -1 else len(text)
            seg_lens.append(seg_end - seg_start)
        if max(seg_lens) - min(seg_lens) <= 1:
            should_bold.add(info["index"])

    return should_bold


def _check_format(info: dict, should_bold: bool = False):
    """检查单个段落的格式是否符合规范：字体/字号/加粗/对齐/缩进/行距"""
    etype = info["element_type"]
    if etype == "empty":
        return

    rules = FONT_RULES.get(etype)
    if not rules:
        return

    current = info["current_format"]
    issues = info["issues"]

    expected_indent = "1.13cm" if rules.get("first_line_indent") else "无"
    info["expected_format"] = {
        "font_name": rules["cn_font"],
        "font_size": str(rules["size"]),
        "bold": "首句加粗" if should_bold else "不加粗",
        "alignment": str(rules.get("alignment")),
        "first_line_indent": expected_indent,
        "line_spacing": "28磅(固定)",
    }

    # ── 字体 ──
    if current["font_name"] and current["font_name"] != rules["cn_font"]:
        if current["font_name"] not in ("Times New Roman", "TimesNewRoman"):
            issues.append({
                "desc": f"字体应为「{rules['cn_font']}」，当前为「{current['font_name']}」",
                "severity": "error",
            })

    # ── 字号 ──
    expected_pt = rules["size"]
    if current["font_size"]:
        current_pt = current["font_size"]
        if current_pt != str(expected_pt):
            issues.append({
                "desc": f"字号应为 {expected_pt}，当前为 {current_pt}",
                "severity": "error",
            })

    # ── 加粗 ──
    if should_bold:
        if not current.get("first_sentence_bold"):
            issues.append({
                "desc": "首句应当加粗（连续等长小标题）",
                "severity": "warning",
            })
    else:
        if current["bold"]:
            issues.append({
                "desc": "不应加粗",
                "severity": "warning",
            })

    # ── 对齐 ──
    expected_align = rules.get("alignment")
    if expected_align is not None:
        cur_align_str = current.get("alignment")
        exp_align_str = str(expected_align)
        if cur_align_str and cur_align_str != exp_align_str:
            exp_name = _ALIGN_NAMES.get(expected_align, exp_align_str)
            cur_name = _ALIGN_NAMES.get(cur_align_str, cur_align_str) if cur_align_str else "无"
            # 对齐名称映射
            align_labels = {
                "CENTER (1)": "居中", "LEFT (0)": "左对齐", "RIGHT (2)": "右对齐",
                "JUSTIFY (3)": "两端对齐", "DISTRIBUTE (4)": "分散对齐",
                "CENTER": "居中", "LEFT": "左对齐", "RIGHT": "右对齐",
                "JUSTIFY": "两端对齐", "DISTRIBUTE": "分散对齐",
            }
            exp_label = align_labels.get(exp_align_str, exp_align_str)
            cur_label = align_labels.get(cur_align_str, cur_align_str)
            if exp_label != cur_label:
                issues.append({
                    "desc": f"对齐方式应为「{exp_label}」，当前为「{cur_label}」",
                    "severity": "warning",
                })

    # ── 首行缩进 ──
    if rules.get("first_line_indent"):
        cur_indent = current.get("first_line_indent", "无")
        if cur_indent == "无" or cur_indent == "0.00cm":
            issues.append({
                "desc": "缺少首行缩进（应为 1.13cm ≈ 2字符）",
                "severity": "error",
            })
    else:
        cur_indent = current.get("first_line_indent", "无")
        if cur_indent not in ("无", "0.00cm", None):
            issues.append({
                "desc": f"不应有首行缩进，当前为 {cur_indent}",
                "severity": "warning",
            })

    # ── 行距 ──
    cur_rule = current.get("line_rule", "?")
    cur_spacing = current.get("line_spacing", "?")
    if cur_rule == 'exact':
        # 固定行距应等于 28pt（允许 ±2pt 误差）
        if cur_spacing != "?":
            try:
                sp_val = float(cur_spacing.replace("磅(固定)", ""))
                if abs(sp_val - 28) > 2:
                    issues.append({
                        "desc": f"固定行距应为 28 磅，当前为 {cur_spacing}",
                        "severity": "error",
                    })
            except ValueError:
                pass
    elif cur_rule == 'auto':
        issues.append({
            "desc": f"行距应为固定值 28 磅，当前为倍数行距（{cur_spacing}）",
            "severity": "error",
        })
    elif cur_rule == '?':
        pass  # 无法读取则不报
    else:
        issues.append({
            "desc": f"行距应为固定值 28 磅，当前行距规则为 {cur_rule}",
            "severity": "error",
        })

    # ── 段前段后间距 ──
    sb = current.get("space_before", "0pt")
    sa = current.get("space_after", "0pt")
    if sb != "0pt":
        issues.append({
            "desc": f"段前间距应为 0，当前为 {sb}",
            "severity": "warning",
        })
    if sa != "0pt":
        issues.append({
            "desc": f"段后间距应为 0，当前为 {sa}",
            "severity": "warning",
        })


def _check_page_setup(doc: Document) -> list:
    """检查文档页面设置：页边距"""
    issues = []
    for section in doc.sections:
        top_mm = section.top_margin / 36000
        bottom_mm = section.bottom_margin / 36000
        left_mm = section.left_margin / 36000
        right_mm = section.right_margin / 36000

        if abs(top_mm - 37) > 1:
            issues.append({
                "desc": f"上边距应为 37mm，当前为 {top_mm:.0f}mm",
                "severity": "error",
            })
        if abs(bottom_mm - 35) > 1:
            issues.append({
                "desc": f"下边距应为 35mm，当前为 {bottom_mm:.0f}mm",
                "severity": "error",
            })
        if abs(left_mm - 28) > 1:
            issues.append({
                "desc": f"左边距应为 28mm，当前为 {left_mm:.0f}mm",
                "severity": "error",
            })
        if abs(right_mm - 26) > 1:
            issues.append({
                "desc": f"右边距应为 26mm，当前为 {right_mm:.0f}mm",
                "severity": "error",
            })
    return issues

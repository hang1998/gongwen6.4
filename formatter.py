"""格式应用引擎 — 页面/段落/字符级格式统一"""

import os
import copy
import re
import uuid
import zipfile
import shutil
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

from config import (
    PAGE_MARGINS, LINE_SPACING, PARA_SPACING_BEFORE, PARA_SPACING_AFTER,
    FONT_RULES,
)
from parser import parse_document
import utils

# 上传文件临时目录
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "outputs")


def _extract_footnote_parts(docx_path: str) -> dict:
    """从原始 docx（ZIP）中提取脚注/尾注 XML 部件（python-docx 可能在保存时丢弃它们）。"""
    parts = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            low = name.lower()
            if ('footnote' in low or 'endnote' in low) and name.endswith('.xml'):
                parts[name] = z.read(name)
    return parts


def _fix_footnote_line_spacing(parts: dict):
    """将脚注/尾注 XML 中所有段落的行距改为单倍行距（lineRule=auto, line=240）。"""
    for name, data in parts.items():
        root = etree.fromstring(data)
        for p in root.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(p, qn('w:pPr'))
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = etree.SubElement(pPr, qn('w:spacing'))
            _clear_spacing_attrs(spacing)
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')
        parts[name] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def _inject_parts(docx_path: str, parts: dict):
    """将保留的部件写入 docx 文件，覆盖已存在的同路径条目。"""
    temp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in parts:
                    zout.writestr(item, parts[item.filename])
                else:
                    zout.writestr(item, zin.read(item.filename))
            # 写入原始 ZIP 中有但输出中丢失的部件（如 word/footnotes.xml）
            for name, data in parts.items():
                if name not in zin.namelist():
                    zout.writestr(name, data)
    shutil.move(temp_path, docx_path)


def format_document(file_id: str) -> str:
    """
    对指定文件执行格式统一，返回修正后文件的 ID。
    原文件只读不写。
    """
    input_path = os.path.join(UPLOAD_DIR, f"{file_id}.docx")
    output_id = str(uuid.uuid4())
    output_path = os.path.join(OUTPUT_DIR, f"{output_id}.docx")

    # 从原始 ZIP 中提取脚注/尾注部件（python-docx 可能在保存时丢弃它们）
    preserved_parts = _extract_footnote_parts(input_path)

    doc = Document(input_path)
    structure = parse_document(input_path)

    # ── 0. 样式表默认间距归零（解决 Word 内置样式覆盖问题）──
    _reset_style_defaults(doc)

    # ── 1. 页面设置 ──
    _apply_page_setup(doc)

    # ── 1b. 全文段前段后间距归零 ──
    _reset_all_paragraph_spacing(doc)

    # ── 1c. 标题破折号拆分（必须在附件预处理前，因为会插入新段落）──
    _handle_title_dash(doc, structure["paragraphs"])

    # ── 1d. 结尾附件预处理：合并首项 → 数字对齐 ──
    _merge_attachment_first_item(doc, structure["paragraphs"])
    _prep_attachment_items(doc, structure["paragraphs"])

    # ── 2. 逐段落格式设置 ──
    is_regulation = structure.get("is_regulation", False)
    for info in structure["paragraphs"]:
        etype = info["element_type"]
        para = doc.paragraphs[info["index"]]

        if etype == "empty":
            if _para_has_image(para):
                _set_single_line_spacing(para)
            continue

        rules = FONT_RULES.get(etype)
        if not rules:
            continue

        _apply_paragraph_format(para, rules, etype)

        # 制度类：第XX章/第XX节 空格规范化 + 第XX条加粗
        if is_regulation:
            if etype == "regulation_chapter":
                _format_regulation_heading(para, re.compile(r'^(第[一二三四五六七八九十百千\d]+章)\s*'), "黑体")
            elif etype == "regulation_section":
                _format_regulation_heading(para, re.compile(r'^(第[一二三四五六七八九十百千\d]+节)\s*'), "楷体_GB2312")
            elif etype == "regulation_article":
                _format_regulation_article(para)

    # ── 2a. 移除段间空行 ──
    _remove_empty_paragraphs(doc, structure["paragraphs"])

    # ── 2c. 落款前插入两个空行（必须在附件 spacing 之前，因为附件 spacing 会改动 body 结构，
    #     导致 paragraphs_info 索引偏移。附件 spacing 插入的新段落不会影响 signature 前空行的计数）──
    _insert_signature_spacing(doc, structure["paragraphs"])

    # ── 2b. 拆分结尾附件中的多编号项（在格式循环后，避免索引偏移）──
    _split_attachment_content(doc, structure["paragraphs"])

    # ── 2c. 结尾附件段前 + 附件项之间插入空行 ──
    _insert_attachment_spacing(doc, structure["paragraphs"])

    # ── 2d. 小标题自动加粗 ──
    _apply_subtitle_bold(doc)

    # ── 2e. 大标题与正文间插入一个空行 ──
    _insert_title_spacing(doc, structure["paragraphs"])

    # ── 2f. 制度类章/节间距处理（必须在 title spacing 之后）──
    if is_regulation:
        _apply_regulation_spacing(doc, structure["paragraphs"])

    # ── 4. 保存 ──
    doc.save(output_path)

    # ── 4b. 将脚注/尾注部件注入回输出文件（先修正行距为单倍行距）──
    if preserved_parts:
        _fix_footnote_line_spacing(preserved_parts)
        _inject_parts(output_path, preserved_parts)

    return output_id


def _apply_page_setup(doc: Document):
    """设置页面：边距、A4纸张"""
    for section in doc.sections:
        section.top_margin = PAGE_MARGINS["top"]
        section.bottom_margin = PAGE_MARGINS["bottom"]
        section.left_margin = PAGE_MARGINS["left"]
        section.right_margin = PAGE_MARGINS["right"]
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def _clear_spacing_attrs(sp_element):
    """彻底清除一个 w:spacing 元素中的所有段间距属性"""
    sp_element.set(qn('w:before'), '0')
    sp_element.set(qn('w:after'), '0')
    # 清除行数单位的间距（Word 中显示为"X行"）
    sp_element.set(qn('w:beforeLines'), '0')
    sp_element.set(qn('w:afterLines'), '0')
    # 清除自动间距
    sp_element.set(qn('w:beforeAutospacing'), '0')
    sp_element.set(qn('w:afterAutospacing'), '0')


def _reset_style_defaults(doc: Document):
    """
    直接修改 styles.xml 中的 docDefaults 和 Normal 样式，
    从根源上消除 Word 内置默认段间距。
    """
    styles_element = doc.styles.element
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    # ── 1. 修改 docDefaults/pPrDefault 确保所有段落的默认段间距为 0 ──
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_element, qn('w:docDefaults'))
    pPrDefault = docDefaults.find(qn('w:pPrDefault'))
    if pPrDefault is None:
        pPrDefault = etree.SubElement(docDefaults, qn('w:pPrDefault'))
    pPr = pPrDefault.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(pPrDefault, qn('w:pPr'))
    sp_default = pPr.find(qn('w:spacing'))
    if sp_default is None:
        sp_default = etree.SubElement(pPr, qn('w:spacing'))
    _clear_spacing_attrs(sp_default)
    sp_default.set(qn('w:line'), str(int(LINE_SPACING.pt * 20)))
    sp_default.set(qn('w:lineRule'), 'exact')

    # ── 2. 修改 Normal 样式（默认段落样式）的段间距为 0 ──
    for style in styles_element.findall(qn('w:style')):
        is_default = style.get(qn('w:default'))
        style_type = style.get(qn('w:type'))
        if is_default == '1' and style_type == 'paragraph':
            style_pPr = style.find(qn('w:pPr'))
            if style_pPr is None:
                style_pPr = etree.SubElement(style, qn('w:pPr'))
            style_sp = style_pPr.find(qn('w:spacing'))
            if style_sp is None:
                style_sp = etree.SubElement(style_pPr, qn('w:spacing'))
            _clear_spacing_attrs(style_sp)
            break


def _reset_all_paragraph_spacing(doc: Document):
    """全文所有段落间距归零，行距统一设为 28 磅固定值（XML 级别，确保写入）"""
    for para in doc.paragraphs:
        pPr = para._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = etree.SubElement(pPr, qn('w:spacing'))
        _clear_spacing_attrs(spacing)
        spacing.set(qn('w:line'), str(int(LINE_SPACING.pt * 20)))
        spacing.set(qn('w:lineRule'), 'exact')
        # 同时通过 python-docx API 设置（双保险）
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


# 段落 Run 内允许通过纯文本重建的安全子元素
# lastRenderedPageBreak: Word 自动插入的分页标记，不影响内容
# del, ins: 修订模式下的删除/插入标记
_SAFE_RUN_CHILDREN = {'rPr', 't', 'br', 'cr', 'tab', 'noBreakHyphen', 'softHyphen', 'sym',
                      'lastRenderedPageBreak', 'del', 'ins'}


def _run_has_protected_content(run) -> bool:
    """检查单个 Run 是否包含不可重建的特殊子元素（脚注/尾注/图片/域代码等）"""
    for child in run._element:
        if etree.QName(child).localname not in _SAFE_RUN_CHILDREN:
            return True
    return False


def _para_has_protected_content(para) -> bool:
    """检查段落中是否有任何 Run 包含不可重建的特殊内容"""
    for run in para.runs:
        if _run_has_protected_content(run):
            return True
    return False


def _para_has_image(para) -> bool:
    """精确检查段落是否包含图片/绘图/嵌入对象（需单倍行距避免裁剪）。"""
    # w:drawing — OOXML 图片/图形
    if para._element.findall('.//' + qn('w:drawing')):
        return True
    # w:pict — 旧版 VML 图片
    if para._element.findall('.//' + qn('w:pict')):
        return True
    # w:object — 嵌入对象（Excel 图表等）
    if para._element.findall('.//' + qn('w:object')):
        return True
    return False


def _clear_paragraph(para):
    """清空段落格式（空段落占位）"""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))
    _clear_spacing_attrs(spacing)
    # 清空内容
    for run in para.runs:
        run.text = ""


def _apply_paragraph_format(para, rules: dict, etype: str = ""):
    """对单个段落应用格式规则"""
    pf = para.paragraph_format

    # 对齐方式
    pf.alignment = rules.get("alignment", WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 首行缩进 — 先清除旧缩进属性（包括 firstLineChars 避免冲突）
    _clear_indent(para)
    indent = rules.get("first_line_indent")
    if indent is not None:
        pf.first_line_indent = indent

    # 段前段后
    pf.space_before = PARA_SPACING_BEFORE
    pf.space_after = PARA_SPACING_AFTER

    # 行距：含图片等特殊内容的段落使用单倍行距，避免图片被裁剪
    if _para_has_image(para):
        _set_single_line_spacing(para)
    else:
        _set_fixed_line_spacing(para, LINE_SPACING)

    # 字符级格式：CJK/ASCII 分割 + 字体设置
    _apply_run_format(para, rules, etype)


def _clear_indent(para):
    """清除段落所有缩进属性，避免 firstLineChars 等残留属性与新值冲突"""
    pPr = para._element.get_or_add_pPr()
    ind_el = pPr.find(qn('w:ind'))
    if ind_el is not None:
        pPr.remove(ind_el)


def _set_fixed_line_spacing(para, spacing_pt):
    """通过 XML 设置固定行距（python-docx 高层 API 不完善）"""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))

    # 行距值单位：exact/atLeast 模式下为 1/20 磅 (ST_TwipsMeasure)
    line_value = int(spacing_pt.pt * 20)
    spacing.set(qn('w:line'), str(line_value))
    spacing.set(qn('w:lineRule'), 'exact')
    _clear_spacing_attrs(spacing)


def _set_single_line_spacing(para):
    """设置单倍行距（用于图片所在段落，避免图片被固定行距裁剪）"""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))
    _clear_spacing_attrs(spacing)
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')


def _apply_run_format(para, rules: dict, etype: str = ""):
    """
    字符级格式设置：
    1. 检查是否包含不可重建的特殊元素（图片/脚注/尾注等），如有则仅更新字体属性不重建 Run
    2. 无特殊元素时：合并文本 → CJK/ASCII 分割 → 重建 Run
    """
    cn_font = rules["cn_font"]
    en_font = rules["en_font"]
    font_size = rules["size"]
    bold = rules["bold"]

    # 检查段落是否包含不可重建的特殊内容
    if _para_has_protected_content(para):
        # 仅更新现有 Run 的字体属性，不删除重建
        for run in para.runs:
            run.font.size = font_size
            run.font.bold = bold
            run.font.name = cn_font
            _set_east_asian_font(run, cn_font)
            _set_ascii_font(run, en_font)
        return

    # 无图片：合并文本，清空重建
    full_text = "".join(run.text for run in para.runs if run.text)
    # 去除段首空白字符；附件编号项保留全角空格用于数字对齐
    if etype in ("attachment_end", "attachment_item"):
        full_text = re.sub(r'^[ \t\r\n\f\v]+', '', full_text)
    else:
        full_text = full_text.lstrip()

    for run in para.runs:
        run._element.getparent().remove(run._element)

    if not full_text.strip():
        return

    # 标题段落：仅将标题格式套用到第一个句号为止，句号之后用正文格式
    _heading_patterns = {
        'heading_1': utils.HEADING_1,
        'heading_2': utils.HEADING_2,
        'heading_3': utils.HEADING_3,
    }
    body_rules = FONT_RULES.get('body')
    if etype in _heading_patterns and body_rules:
        pm = _heading_patterns[etype].match(full_text)
        if pm:
            content = full_text[pm.end():]
            dot_pos = content.find('。')
            if dot_pos != -1 and dot_pos < len(content) - 1:
                split_pos = pm.end() + dot_pos + 1
                _add_formatted_runs(para, full_text[:split_pos], cn_font, en_font, font_size, bold)
                _add_formatted_runs(para, full_text[split_pos:], body_rules["cn_font"], body_rules["en_font"], body_rules["size"], body_rules["bold"])
                return

    _add_formatted_runs(para, full_text, cn_font, en_font, font_size, bold)


def _add_formatted_runs(para, text, cn_font, en_font, font_size, bold):
    """按 CJK/ASCII 分割文本并创建格式化 Run。"""
    segments = utils.split_cjk_ascii(text)
    for seg_text, is_cjk in segments:
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        run.font.size = font_size
        run.font.bold = bold
        if is_cjk:
            run.font.name = cn_font
            _set_east_asian_font(run, cn_font)
        else:
            run.font.name = en_font
            _set_east_asian_font(run, cn_font)
        _set_ascii_font(run, en_font)


def _set_east_asian_font(run, font_name: str):
    """设置东亚字体（w:eastAsia）"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_ascii_font(run, font_name: str):
    """设置西文字体（w:ascii + w:hAnsi）"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def _apply_subtitle_bold(doc: Document):
    """自动检测小标题并加粗：
    1. 连续3段以上正文，首句（到第一个。）字数相同 → 每段首句加粗
    2. 段内出现3次以上"一是/二是/XX是"且各段到。字数相同 → 加粗这些分句（到。）
    """
    import re
    body_paras = [p for p in doc.paragraphs if p.text.strip()]

    # ── 模式1：检测连续段落首句长度相同的分组 ──
    first_sentences = []
    for p in body_paras:
        text = p.text.strip()
        dot_pos = text.find('。')
        if dot_pos == -1:
            first_sentences.append(None)
        else:
            first_sentences.append((p, dot_pos + 1, len(text[:dot_pos])))

    i = 0
    while i < len(first_sentences):
        if first_sentences[i] is None:
            i += 1
            continue
        _, _, length = first_sentences[i]
        j = i + 1
        while j < len(first_sentences):
            if first_sentences[j] is None or first_sentences[j][2] != length:
                break
            j += 1
        if j - i >= 3:
            for k in range(i, j):
                para, end_pos, _ = first_sentences[k]
                _bold_char_range(para, 0, end_pos)
            i = j
        else:
            i += 1

    # ── 模式2：段内"一是/二是/三是..."等分句加粗（到。）──
    branch_re = re.compile(r'[一二三四五六七八九十]+[、，,是]|\d+[、，,\.是]')
    for p in body_paras:
        text = p.text.strip()
        # 找出所有"XX是"或"XX、"标记的位置
        markers = []
        for m in branch_re.finditer(text):
            # 确保在句首或分句开头位置
            pos = m.start()
            if pos == 0 or text[pos-1] in '。；;，,、\n':
                markers.append((m.start(), m.end()))
        if len(markers) < 3:
            continue
        # 每段：从标记开头到下一个。或段落末尾
        segments = []
        for idx, (seg_start, _) in enumerate(markers):
            # 找这个分句的结束位置（下一个。或段落末尾）
            dot = text.find('。', seg_start)
            if dot == -1:
                seg_end = len(text)
            else:
                seg_end = dot + 1  # 包含。
            segments.append((seg_start, seg_end))
        if len(segments) < 3:
            continue
        # 检查各段字符数是否相同（±1）
        seg_lens = [end - start for start, end in segments]
        if max(seg_lens) - min(seg_lens) <= 1:
            for seg_start, seg_end in segments:
                _bold_char_range(p, seg_start, seg_end)


def _bold_char_range(para, start_char: int, end_char: int):
    """在段落中按字符位置精确加粗 [start_char, end_char) 范围。
    会拆分跨边界的 run，确保只加粗目标范围内的字符。"""
    if start_char >= end_char:
        return
    # 收集每个 run 的文本和字符范围
    runs_info = []
    pos = 0
    for run in para.runs:
        t = run.text if run.text else ''
        runs_info.append((run, t, pos, pos + len(t)))
        pos += len(t)

    # 找到 start_char 和 end_char 所在的 run 索引
    start_run_idx = end_run_idx = -1
    for idx, (_, _, s, e) in enumerate(runs_info):
        if start_run_idx == -1 and start_char < e:
            start_run_idx = idx
        if end_run_idx == -1 and end_char <= e:
            end_run_idx = idx
        if start_run_idx >= 0 and end_run_idx >= 0:
            break

    if start_run_idx < 0 or end_run_idx < 0:
        return

    # 从后往前处理，避免索引变化
    for idx in range(end_run_idx, start_run_idx - 1, -1):
        run, text, s, e = runs_info[idx]
        # 含有脚注/尾注等受保护内容的 Run 只整体加粗，不拆分
        safe_to_split = not _run_has_protected_content(run)

        if idx == start_run_idx and idx == end_run_idx:
            # 起止在同一个 run 内：拆成 [s, start) [start, end) [end, e)，加粗中间
            if safe_to_split and end_char < e:
                _split_run_at(para, run, end_char - s)
            if safe_to_split and start_char > s:
                after_start = _split_run_at(para, run, start_char - s)
                _set_bold(after_start)
            else:
                _set_bold(run)
        elif idx == start_run_idx:
            if safe_to_split and start_char > s:
                new_run = _split_run_at(para, run, start_char - s)
                _set_bold(new_run)
            else:
                _set_bold(run)
        elif idx == end_run_idx:
            if safe_to_split and end_char < e:
                _split_run_at(para, run, end_char - s)
            _set_bold(run)
        else:
            _set_bold(run)


def _split_run_at(para, run, offset: int):
    """在 run 内 offset 字符处拆分，返回后半部分的新 Run。
    使用 para.add_run() 确保创建的是正确的 python-docx Run 对象。"""
    text = run.text if run.text else ''
    if offset <= 0 or offset >= len(text):
        return run
    left_text = text[:offset]
    right_text = text[offset:]

    # 保存原 run 的格式
    saved_name = run.font.name
    saved_size = run.font.size
    saved_bold = run.font.bold
    from docx.oxml.ns import qn as _qn
    rPr_orig = run._element.find(_qn('w:rPr'))

    # 用 para.add_run 创建新 run（python-docx 正确包装）
    new_run = para.add_run(right_text)
    if rPr_orig is not None:
        # 复制原格式属性到新 run
        new_rPr = new_run._element.find(_qn('w:rPr'))
        if new_rPr is not None:
            new_run._element.remove(new_rPr)
        new_run._element.insert(0, etree.fromstring(etree.tostring(rPr_orig)))
    new_run.font.bold = saved_bold

    # 修改原 run 文本为前半部分
    run.text = left_text
    _set_bold(run) if saved_bold else None

    # 将新 run 移动到原 run 之后
    run._element.addnext(new_run._element)

    return new_run


def _set_bold(run):
    """设置 run 为加粗"""
    run.font.bold = True


def _handle_title_dash(doc, paragraphs_info):
    """若标题中包含 —— 或 —，在破折号处拆分为标题 + 副标题（楷体）。"""
    for i, info in enumerate(paragraphs_info):
        if info["element_type"] != "title":
            continue
        text = info["text"]
        match = re.search(r'(——)|(—)', text)
        if not match:
            return
        pos = match.start()
        if pos == 0:
            info["element_type"] = "subtitle"
            return

        before = text[:pos].strip()
        after = text[pos:].strip()

        para = doc.paragraphs[info["index"]]
        para.text = before
        info["text"] = before

        sub_elem = _make_subtitle_paragraph(after)
        para._element.addnext(sub_elem)

        subtitle_info = {
            "index": info["index"] + 1,
            "text": after,
            "element_type": "subtitle",
            "issues": [],
            "current_format": {
                "font_name": None, "font_size": None, "bold": None,
                "alignment": None, "first_line_indent": "无",
                "line_spacing": "?", "line_rule": "?",
                "space_before": "0pt", "space_after": "0pt",
                "first_sentence_bold": False,
            },
            "expected_format": None,
        }
        paragraphs_info.insert(i + 1, subtitle_info)
        for j in range(i + 2, len(paragraphs_info)):
            paragraphs_info[j]["index"] += 1
        return


def _make_subtitle_paragraph(text):
    """创建副标题段落元素（楷体_GB2312, 16pt, 居中）。"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'center')
    sp = etree.SubElement(pPr, qn('w:spacing'))
    _clear_spacing_attrs(sp)
    sp.set(qn('w:line'), str(int(LINE_SPACING.pt * 20)))
    sp.set(qn('w:lineRule'), 'exact')
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '32')
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '32')
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _remove_empty_paragraphs(doc, paragraphs_info):
    """移除文档中所有无内容的空段落（保护含图片等特殊内容的段落），并更新索引。"""
    empty_info = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip() if para.text else ""
        if not text and not _para_has_image(para):
            empty_info.append((i, para._element))

    if not empty_info:
        return

    body = doc.element.body
    for _, elem in empty_info:
        body.remove(elem)

    empty_indices = [idx for idx, _ in empty_info]
    for info in paragraphs_info:
        shift = sum(1 for idx in empty_indices if idx < info["index"])
        info["index"] -= shift


# ── 结尾附件处理 ──


def _get_attachment_info(paragraphs_info):
    """扫描 parser 输出，返回结尾附件段的结构信息。

    返回:
      (end_idx, title_is_numbered, item_indices)
        - end_idx: 第一个 attachment_end 的 index，无则 None
        - title_is_numbered: attachment_end 文本中是否已包含编号（如"附件：1.xx"）
        - item_indices: 所有有实质内容的 attachment_item 段落 index 列表
    """
    import re
    end_idx = None
    item_indices = []

    for info in paragraphs_info:
        if end_idx is None and info["element_type"] == "attachment_end":
            end_idx = info["index"]
        if info["element_type"] == "attachment_item" and info["text"].strip():
            item_indices.append(info["index"])

    if end_idx is not None:
        title_text = paragraphs_info[end_idx]["text"]
        # 检查标题中是否包含编号（不在行首也可能，如"附件：1.xxx"）
        title_is_numbered = bool(re.search(r'\d+[\.\、．]', title_text))
    else:
        title_is_numbered = False

    return end_idx, title_is_numbered, item_indices


FULLWIDTH_SPACE = "　"  # 全角空格，与中文字符等宽


def _prep_attachment_items(doc, paragraphs_info):
    """在格式循环前：为结尾附件的编号项前加 3 个全角空格，使数字对齐。

    对齐逻辑：
      - 若 attachment_end 已含编号（如"附件：1.xx"），所有 attachment_item 都是第 2+ 项，全加空格
      - 若 attachment_end 不含编号（如"附件："），第一个 attachment_item 是第 1 项不加空格，其余加
    """
    end_idx, title_is_numbered, item_indices = _get_attachment_info(paragraphs_info)
    if end_idx is None:
        return

    for i, idx in enumerate(item_indices):
        need_prefix = title_is_numbered or (i > 0)
        if need_prefix:
            para = doc.paragraphs[idx]
            original = para.text if para.text else ""
            if original.strip():
                para.text = FULLWIDTH_SPACE * 3 + original


def _merge_attachment_first_item(doc, paragraphs_info):
    """若 attachment_end 不含编号（如"附件："），将第一个 attachment_item 拼接到标题行。

    处理边界情况：文档中"附件："单独一行，下一行是"1.项目A"的场景。
    """
    end_idx, title_is_numbered, item_indices = _get_attachment_info(paragraphs_info)
    if end_idx is None or title_is_numbered or not item_indices:
        return

    title_para = doc.paragraphs[end_idx]
    first_item_para = doc.paragraphs[item_indices[0]]

    title_text = (title_para.text or "").strip()
    item_text = (first_item_para.text or "").strip()

    if title_text and item_text:
        merged = title_text + item_text
        title_para.text = merged
        first_item_para.text = ""
        # 同步更新 paragraphs_info
        removed_idx = item_indices[0]
        paragraphs_info[end_idx]["text"] = merged
        # 从 XML body 中移除被合并的段落，避免残留空行
        first_item_para._element.getparent().remove(first_item_para._element)
        # 从 paragraphs_info 中彻底删除该条目，避免格式循环误清空新段落
        paragraphs_info.pop(removed_idx)
        # 更新所有大于 removed_idx 的 index（doc.paragraphs 会重新索引）
        for info in paragraphs_info:
            if info["index"] > removed_idx:
                info["index"] -= 1


def _split_attachment_content(doc, paragraphs_info):
    """将结尾附件中包含多个编号项的段落拆分为独立段落。

    处理场景：段落文本为"1.项目A。2.项目B。3.项目C。"时，
    拆分为三个独立段落，数字自动对齐。
    """
    import re
    # 匹配编号项：1.xxx 或 1、xxx 或 1．xxx
    item_re = re.compile(r'(\d+)[\.\、．]')

    end_idx, title_is_numbered, item_indices = _get_attachment_info(paragraphs_info)

    # 收集所有要检查的段落索引（attachment_end + 实际有内容的段落）
    check_indices = []
    if end_idx is not None:
        check_indices.append(end_idx)
    # 也要检查 item_indices 中的段落（即使 paragraphs_info 中已清空，doc 中可能还有旧内容）
    for idx in item_indices:
        if idx not in check_indices:
            check_indices.append(idx)
    # 额外扫描：任何紧跟在 end_idx 之后且有编号内容的段落
    if end_idx is not None:
        for info in paragraphs_info:
            if info["index"] > end_idx and info["text"].strip():
                if utils.ATTACHMENT_ITEM.search(info["text"]):
                    if info["index"] not in check_indices:
                        check_indices.append(info["index"])
                        t = info["text"][:40]

    body_elem = doc.element.body

    # 从后往前处理，避免插入新段落后索引偏移
    for idx in sorted(check_indices, reverse=True):
        para = doc.paragraphs[idx]
        text = (para.text or "").strip()
        if not text:
            continue

        # 找所有编号项的位置
        items_info = []  # [(number_str, start_pos, end_pos), ...]
        for m in item_re.finditer(text):
            num_str = m.group(1)
            start = m.start()
            # 找下一个编号的起始位置作为当前项的结束
            next_m = item_re.search(text, m.end())
            end = next_m.start() if next_m else len(text)
            if end > start:
                items_info.append((num_str, start, end))


        # 少于 2 个编号项则无需拆分
        if len(items_info) < 2:
            continue

        # 检查是否已经是正确的编号序列（如 1, 2, 3...）
        expected = 1
        all_sequential = True
        for num_str, _, _ in items_info:
            if num_str != str(expected):
                all_sequential = False
                break
            expected += 1
        if not all_sequential:
            continue

        # 不需要拆分：只有"附件："开头但只有1个编号，或没有编号
        # 需要拆分：多个连续的 1.xxx 2.xxx 3.xxx
        if len(items_info) < 2:
            continue

        # 提取各段文本
        segments = []
        for i, (num_str, start, end) in enumerate(items_info):
            seg_text = text[start:end].strip()
            segments.append(seg_text)

        # 如果当前段落是 attachment_end：第一段保留在原段落，其余插入后面
        is_end_para = (idx == end_idx)

        if is_end_para:
            first_seg = segments[0]
            if not first_seg.startswith('附件') and idx == end_idx:
                first_seg = '附件：' + first_seg
            para.text = first_seg
            # 其余项直接另起一行，数字对齐，不加空行
            for seg_text in reversed(segments[1:]):
                item_p = _create_attachment_item_paragraph(FULLWIDTH_SPACE * 3 + seg_text)
                para._element.addnext(item_p)
        else:
            para.text = segments[0]
            for seg_text in reversed(segments[1:]):
                item_p = _create_attachment_item_paragraph(FULLWIDTH_SPACE * 3 + seg_text)
                para._element.addnext(item_p)


def _create_attachment_item_paragraph(text):
    """创建一个 attachment_item 段落元素，含正确格式（仿宋、1.13cm 缩进、固定行距）"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    # 两端对齐
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'both')
    # 首行缩进 1.13cm = 641 twips (1cm = 567 twips, 1.13cm ≈ 641)
    ind = etree.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:firstLine'), '641')
    # 固定行距 28pt = 560 (28*20)
    sp = etree.SubElement(pPr, qn('w:spacing'))
    _clear_spacing_attrs(sp)
    sp.set(qn('w:line'), str(int(LINE_SPACING.pt * 20)))
    sp.set(qn('w:lineRule'), 'exact')
    # 文本内容
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '32')  # 16pt = 32 half-pts
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '32')
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _insert_attachment_spacing(doc, paragraphs_info):
    """在结尾附件段前插入 1 个空行（与正文分隔）。项目之间不插入空行。"""
    end_idx, _, item_indices = _get_attachment_info(paragraphs_info)

    # 收集所有结尾附件相关段落的 index
    attach_indices = []
    if end_idx is not None:
        attach_indices.append(end_idx)
    attach_indices.extend(item_indices)

    # 额外扫描 end_idx 之后有全角空格前缀的段落（_split_attachment_content 创建的）
    if end_idx is not None:
        for i in range(end_idx + 1, len(doc.paragraphs)):
            text = (doc.paragraphs[i].text or '').strip()
            if not text:
                continue
            if text.startswith(FULLWIDTH_SPACE):
                if i not in attach_indices:
                    attach_indices.append(i)
            else:
                break

    if not attach_indices:
        return
    attach_indices.sort()

    body_elem = doc.element.body

    # 附件段前插入 1 个空行
    first_para = doc.paragraphs[attach_indices[0]]
    first_elem = first_para._element
    _ensure_blank_before(body_elem, first_elem, 1)


def _ensure_blank_before(body_elem, target_elem, min_blanks):
    """确保 target_elem 前面至少有 min_blanks 个空段落"""
    body_children = list(body_elem)
    pos = None
    for i, child in enumerate(body_children):
        if child is target_elem:
            pos = i
            break
    if pos is None:
        return
    existing = 0
    for j in range(pos - 1, -1, -1):
        child = body_children[j]
        if child.tag == qn('w:p'):
            texts = child.findall('.//' + qn('w:t'))
            if not any(t.text and t.text.strip() for t in texts):
                existing += 1
                continue
        break
    need = max(0, min_blanks - existing)
    for _ in range(need):
        empty_p = _make_empty_paragraph()
        target_elem.addprevious(empty_p)


def _insert_blank_between(body_elem, elem_a, elem_b):
    """在 elem_a 和 elem_b 之间插入一个空段落（如果还没有）"""
    body_children = list(body_elem)
    pos_a = pos_b = None
    for i, child in enumerate(body_children):
        if child is elem_a:
            pos_a = i
        if child is elem_b:
            pos_b = i
    if pos_a is None or pos_b is None:
        return
    # 检查之间是否已有非空内容或空段落
    has_gap = False
    for j in range(pos_a + 1, pos_b):
        child = body_children[j]
        if child.tag == qn('w:p'):
            texts = child.findall('.//' + qn('w:t'))
            if not any(t.text and t.text.strip() for t in texts):
                has_gap = True  # 已有空段落
                break
        # 非 p 元素视为内容，已有间隔
        has_gap = True
        break
    if not has_gap:
        empty_p = _make_empty_paragraph()
        elem_b.addprevious(empty_p)


def _insert_title_spacing(doc, paragraphs_info):
    """在大标题/副标题区域之后插入一个空行，与正文分隔。"""
    last_title_idx = None
    for info in paragraphs_info:
        if info["element_type"] in ("title", "subtitle"):
            last_title_idx = info["index"]
    if last_title_idx is None:
        return

    body_elem = doc.element.body
    title_elem = doc.paragraphs[last_title_idx]._element
    empty_p = _make_empty_paragraph()
    title_elem.addnext(empty_p)


def _make_empty_paragraph():
    """创建一个带零间距+固定行距的空段落元素"""
    empty_p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(empty_p, qn('w:pPr'))
    sp = etree.SubElement(pPr, qn('w:spacing'))
    _clear_spacing_attrs(sp)
    sp.set(qn('w:line'), str(int(LINE_SPACING.pt * 20)))
    sp.set(qn('w:lineRule'), 'exact')
    return empty_p


def _insert_signature_spacing(doc: Document, paragraphs_info: list):
    """在落款前插入两个空行（若前面已有空行则补足到两个）"""
    # 找到第一个签名或日期段落
    sig_idx = None
    for info in paragraphs_info:
        if info["element_type"] in ("signature", "signature_date"):
            sig_idx = info["index"]
            break
    if sig_idx is None:
        return

    # 统计签名段前已有的连续空段落数（从签名段向前数）
    body = doc.element.body
    body_children = list(body)
    # 找到签名段在 body 子元素中的位置
    sig_para = doc.paragraphs[sig_idx]
    sig_elem = sig_para._element
    sig_pos = None
    for i, child in enumerate(body_children):
        if child is sig_elem:
            sig_pos = i
            break
    if sig_pos is None:
        return

    # 向前统计已有空段落
    existing_empty = 0
    for j in range(sig_pos - 1, -1, -1):
        child = body_children[j]
        if child.tag == qn('w:p'):
            # 检查是否为空段落
            pPr = child.find(qn('w:pPr'))
            texts = child.findall('.//' + qn('w:t'))
            has_text = any(t.text and t.text.strip() for t in texts)
            if not has_text:
                existing_empty += 1
                continue
        break

    # 补足到两个空行
    need = max(0, 2 - existing_empty)
    for _ in range(need):
        empty_p = etree.SubElement(body, qn('w:p'))
        # 设置空行段间距为 0 + 固定行距 28 磅
        pPr = etree.SubElement(empty_p, qn('w:pPr'))
        sp = etree.SubElement(pPr, qn('w:spacing'))
        _clear_spacing_attrs(sp)
        sp.set(qn('w:line'), str(int(LINE_SPACING.pt * 20)))
        sp.set(qn('w:lineRule'), 'exact')
        # 在签名段之前插入
        body.remove(empty_p)
        sig_elem.addprevious(empty_p)


# ── 制度类格式化 ──

# 匹配"第XX条"前缀（含中文数字和阿拉伯数字）
_REGULATION_ARTICLE_PREFIX = re.compile(r'^(第[一二三四五六七八九十百千\d]+条)\s*')


def _format_regulation_article(para):
    """格式化制度类条文：'第XX条'加粗，后接一个空格，删除多余空格"""
    full_text = "".join(run.text for run in para.runs if run.text)
    if not full_text.strip():
        return

    # 去除段首空白
    full_text = full_text.lstrip()

    m = _REGULATION_ARTICLE_PREFIX.match(full_text)
    if not m:
        return

    prefix = m.group(1)  # 如 "第三条"
    rest = full_text[m.end():]  # 剩余文本（可能以空格开头）
    # 规范化：删除多余空格，确保"第XX条"后空一格
    rest = rest.lstrip()

    # 清除旧 runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # 加粗的"第XX条" + 空格
    bold_run = para.add_run(prefix + " ")
    bold_run.font.bold = True
    bold_run.font.size = Pt(16)
    bold_run.font.name = "仿宋_GB2312"
    _set_east_asian_font(bold_run, "仿宋_GB2312")
    _set_ascii_font(bold_run, "Times New Roman")

    # 剩余文本按 CJK/ASCII 分割
    if rest:
        segments = utils.split_cjk_ascii(rest)
        for seg_text, is_cjk in segments:
            if not seg_text:
                continue
            run = para.add_run(seg_text)
            run.font.size = Pt(16)
            run.font.bold = False
            if is_cjk:
                run.font.name = "仿宋_GB2312"
                _set_east_asian_font(run, "仿宋_GB2312")
            else:
                run.font.name = "Times New Roman"
                _set_east_asian_font(run, "仿宋_GB2312")
            _set_ascii_font(run, "Times New Roman")


def _format_regulation_heading(para, prefix_re, cn_font):
    """格式化制度类章/节标题：第X章/第X节后空一格接文字，删除多余空格"""
    full_text = "".join(run.text for run in para.runs if run.text)
    if not full_text.strip():
        return

    full_text = full_text.lstrip()
    m = prefix_re.match(full_text)
    if not m:
        return

    prefix = m.group(1)  # 如 "第一章" 或 "第一节"
    rest = full_text[m.end():].lstrip()
    normalized = prefix + " " + rest if rest else prefix

    # 清除旧 runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # 重建格式化 runs
    segments = utils.split_cjk_ascii(normalized)
    for seg_text, is_cjk in segments:
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        run.font.size = Pt(16)
        run.font.bold = False
        if is_cjk:
            run.font.name = cn_font
            _set_east_asian_font(run, cn_font)
        else:
            run.font.name = "Times New Roman"
            _set_east_asian_font(run, cn_font)
        _set_ascii_font(run, "Times New Roman")


def _apply_regulation_spacing(doc, paragraphs_info):
    """制度类章/节间距处理：
    - 第XX章/第XX节：与上一个文字段落之间确保一个空行
    - 第XX章/第XX节：与下一个文字段落之间删除所有空行

    直接扫描 doc.paragraphs，不依赖可能过时的 paragraphs_info 索引。
    """
    import re
    chapter_re = re.compile(r'^\s*第[一二三四五六七八九十百千\d]+章\s')
    section_re = re.compile(r'^\s*第[一二三四五六七八九十百千\d]+节\s')

    body_elem = doc.element.body

    # 从后往前处理，避免元素位置变化导致的问题
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[idx]
        text = (para.text or "").strip()
        if not text:
            continue

        is_chapter = bool(chapter_re.match(text))
        is_section = bool(section_re.match(text))
        if not (is_chapter or is_section):
            continue

        elem = para._element
        body_children = list(body_elem)
        pos = None
        for i, child in enumerate(body_children):
            if child is elem:
                pos = i
                break
        if pos is None:
            continue

        # 1. 删除章/节之后的所有空段落（直到下一个非空段或文档末尾）
        body_children = list(body_elem)
        for j in range(pos + 1, len(body_children)):
            child = body_children[j]
            if child.tag == qn('w:p'):
                texts = child.findall('.//' + qn('w:t'))
                has_text = any(t.text and t.text.strip() for t in texts)
                if not has_text:
                    body_elem.remove(child)
                    body_children = list(body_elem)
                else:
                    break
            else:
                break

        # 刷新位置
        body_children = list(body_elem)
        pos = None
        for i, child in enumerate(body_children):
            if child is elem:
                pos = i
                break
        if pos is None:
            continue

        # 2. 确保章/节之前恰好有一个空行
        if pos == 0:
            continue

        existing_empty = 0
        for j in range(pos - 1, -1, -1):
            child = body_children[j]
            if child.tag == qn('w:p'):
                texts = child.findall('.//' + qn('w:t'))
                has_text = any(t.text and t.text.strip() for t in texts)
                if not has_text:
                    existing_empty += 1
                    continue
            break

        if existing_empty > 1:
            # 删除多余空段落（保留最靠近章/节的 1 个）
            removed = 0
            for j in range(pos - 1, -1, -1):
                child = body_children[j]
                if child.tag == qn('w:p'):
                    texts = child.findall('.//' + qn('w:t'))
                    has_text = any(t.text and t.text.strip() for t in texts)
                    if not has_text:
                        if removed < existing_empty - 1:
                            body_elem.remove(child)
                            removed += 1
                        else:
                            break
                    else:
                        break
                else:
                    break
        elif existing_empty == 0:
            empty_p = _make_empty_paragraph()
            elem.addprevious(empty_p)


